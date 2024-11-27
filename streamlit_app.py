import streamlit as st
import pandas as pd
import folium
from streamlit_folium import st_folium
import h3
from folium.plugins import MarkerCluster
from io import BytesIO
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from docx.shared import Inches



st.title("MEMS 메타데이터 관리")
st.write("SKT MEMS 시각화 및 필터, 메타데이터 열람기능 test ")

latest_status_file = 'Sensor_data_1024.csv'

# 작업 모드 선택
mode = st.radio("작업을 선택하세요:", ("최신 현황 불러오기 ('24년10월 기준)", "새 파일 업로드"))

def format_terminal_number(data):
    if '단말번호' in data.columns:
        data['단말번호'] = data['단말번호'].astype(str).str.zfill(11)
    return data

# 데이터 파일 불러오기
if mode == "새 파일 업로드":
    uploaded_file = st.file_uploader("새 파일을 업로드하세요 (CSV 형식)", type="csv")
    if uploaded_file:
        data = pd.read_csv(uploaded_file, dtype={'단말번호': str})
        data = format_terminal_number(data)
        st.write("새 파일 데이터 미리보기:")
        st.dataframe(data)
else:
    data = pd.read_csv(latest_status_file, dtype={'단말번호': str})
    data = format_terminal_number(data)
    st.write("최신 현황 데이터 미리보기:")
    st.dataframe(data)

# 위도와 경도 열 이름 변경
if '위도' in data.columns and '경도' in data.columns:
    data = data.rename(columns={'위도': 'latitude', '경도': 'longitude'})

# 세션 상태 초기화
if 'show_facility_selection' not in st.session_state:
    st.session_state['show_facility_selection'] = False

# 시설구분(코드명)으로 선택하기 버튼
if st.button("시설구분(코드명)으로 선택하기"):
    st.session_state['show_facility_selection'] = not st.session_state['show_facility_selection']

# 주소 선택 기능 활성화 여부 관리 (세션 상태)
if 'show_address_selection' not in st.session_state:
    st.session_state['show_address_selection'] = False

# "주소로 선택하기" 버튼 클릭 시 상태 토글
if st.button("주소로 선택하기"):
    st.session_state['show_address_selection'] = not st.session_state['show_address_selection']

# 연결상태 선택 기능 활성화 여부 관리 (세션 상태)
if 'show_status_selection' not in st.session_state:
    st.session_state['show_status_selection'] = False

# "연결상태로 선택하기" 버튼 클릭 시 상태 토글
if st.button("연결상태로 선택하기"):
    st.session_state['show_status_selection'] = not st.session_state['show_status_selection']

# 세션 상태 초기화
if 'show_axis_correction_selection' not in st.session_state:
    st.session_state['show_axis_correction_selection'] = False

# "축보정 선택하기" 버튼 생성
if st.button("축보정 선택하기"):
    st.session_state['show_axis_correction_selection'] = not st.session_state['show_axis_correction_selection']

# 필터링 조건 초기화
filtered_data = data.copy()

# 시설구분 선택 활성화 상태에 따른 동작
if st.session_state['show_facility_selection']:
    if '시설구분' in data.columns:
        # 시설구분 체크박스 항목 정의
        facility_options = {
            'SKM': 'SKM(SKT인프라)',
            'POM': 'POM(우정사업본부)',
            'KSM': 'KSM(기상청 관측소)',
            'FSM': 'FSM(소방청)',
            'CPM': 'CPM(해양경찰청)',
        }
        
        selected_facilities = [key for key, label in facility_options.items() if st.checkbox(label, key=f"facility_{key}")]
        
        # SKM 선택 시 시설구분세부 열 참고하여 세부 선택 옵션 표시
        selected_skm_details = []  # 빈 리스트로 초기화
        if 'SKM' in selected_facilities and '시설구분세부' in data.columns:
            st.write("시설구분 세부 선택 (SKM):")
            skm_details = sorted(data[data['시설구분'] == 'SKM']['시설구분세부'].unique())
            select_all_skm_details = st.checkbox("시설구분 세부 전체 선택 (SKM)", key="select_all_skm_details")
            selected_skm_details = [detail for detail in skm_details if st.checkbox(detail, key=f"skm_detail_{detail}", value=select_all_skm_details)]
        
        # 필터링 조건 적용
        facility_filtered_data = data[data['시설구분'].isin(selected_facilities)]
        if selected_skm_details:
            facility_filtered_data = facility_filtered_data[facility_filtered_data['시설구분세부'].isin(selected_skm_details)]
        
        # 필터링 결과를 전체 필터링 데이터에 적용
        filtered_data = filtered_data[filtered_data.index.isin(facility_filtered_data.index)]
    else:
        st.write("데이터에 '시설구분' 열이 없습니다.")

# 주소 선택 기능 활성화 상태에 따른 동작
if st.session_state['show_address_selection']:
    if '주소' in data.columns:
        data['주소_첫단어'] = data['주소'].apply(lambda x: str(x).split()[0] if pd.notna(x) else '')
        data['주소_두번째단어'] = data['주소'].apply(lambda x: str(x).split()[1] if pd.notna(x) and len(str(x).split()) > 1 else '')

        st.write("첫 번째 주소 선택:")
        
        # 첫 번째 주소 전체 선택 기능
        col1, col2 = st.columns(2)
        with col1:
            select_all_first_words = st.checkbox("첫 번째 주소 전체 선택")
            selected_first_words = []
            for word in sorted(data['주소_첫단어'].unique()):
                if st.checkbox(word, key=f"first_{word}", value=select_all_first_words):
                    selected_first_words.append(word)

        # 두 번째 주소 선택 체크박스 생성
        selected_second_words = {}
        with col2:
            for first_word in selected_first_words:
                st.write(f"두 번째 주소 선택 ({first_word}):")
                second_words = sorted(data[data['주소_첫단어'] == first_word]['주소_두번째단어'].unique())

                # 두 번째 주소 전체 선택 기능
                select_all_second_words = st.checkbox(f"{first_word} 전체 선택", key=f"{first_word}_select_all")
                selected_second_words[first_word] = [
                    second_word for second_word in second_words if st.checkbox(second_word, key=f"{first_word}_{second_word}", value=select_all_second_words)
                ]

        # 주소 필터링 조건 적용
        if selected_second_words:
            address_filtered_data = pd.DataFrame()
            for first_word, second_words in selected_second_words.items():
                if second_words:
                    temp_data = data[(data['주소_첫단어'] == first_word) & (data['주소_두번째단어'].isin(second_words))]
                    address_filtered_data = pd.concat([address_filtered_data, temp_data])
            # 전체 필터링 데이터에 주소 필터링 적용
            if not address_filtered_data.empty:
                filtered_data = filtered_data[filtered_data.index.isin(address_filtered_data.index)]
            else:
                filtered_data = pd.DataFrame()  # 조건에 맞는 데이터가 없을 때 빈 데이터프레임으로 설정


# 연결 상태 선택 기능 활성화 상태에 따른 동작
if st.session_state['show_status_selection']:
    if '연결상태' in data.columns:
        st.write("연결 상태 선택:")
        unique_statuses = sorted(data['연결상태'].unique())
        selected_statuses = [status for status in unique_statuses if st.checkbox(status, key=f"status_{status}")]

        # 연결 상태 필터링 조건 적용
        if selected_statuses:
            filtered_data = filtered_data[filtered_data['연결상태'].isin(selected_statuses)]
    else:
        st.write("데이터에 '연결상태' 열이 없습니다.")

# 축보정 선택 기능 활성화 상태에 따른 동작
if st.session_state['show_axis_correction_selection']:
    if '축보정' in data.columns:
        st.write("축보정 선택:")
        
        # 축보정 열의 고유 값을 기반으로 체크박스 생성
        unique_axis_corrections = sorted(data['축보정'].unique())
        selected_axis_corrections = [correction for correction in unique_axis_corrections if st.checkbox(str(correction), key=f"axis_correction_{correction}")]
        
        # 축보정 필터링 조건 적용
        if selected_axis_corrections:
            axis_correction_filtered_data = data[data['축보정'].isin(selected_axis_corrections)]
            # 필터링 결과를 전체 필터링 데이터에 적용
            filtered_data = filtered_data[filtered_data.index.isin(axis_correction_filtered_data.index)]
    else:
        st.write("데이터에 '축보정' 열이 없습니다.")

# 최종 필터링된 데이터 표시
st.write("선택한 조건에 따른 최종 데이터:")
if not filtered_data.empty:
    st.dataframe(filtered_data)
else:
    st.write("선택한 조건에 해당하는 데이터가 없습니다.")


# 지도보기 기능 상태 관리 (세션 상태)
if 'show_map' not in st.session_state:
    st.session_state['show_map'] = False

# "지도보기" 버튼 클릭 시 상태 토글
if st.button("지도보기"):
    st.session_state['show_map'] = not st.session_state['show_map']

# 지도 표시 로직
if st.session_state['show_map']:
    if not filtered_data.empty and 'latitude' in filtered_data.columns and 'longitude' in filtered_data.columns:
        # Folium 지도 생성
        m = folium.Map(location=[filtered_data['latitude'].mean(), filtered_data['longitude'].mean()], zoom_start=10)

        # H3 클러스터링 및 센서 상태 카운트 집계
        filtered_data['h3_index'] = filtered_data.apply(lambda row: h3.latlng_to_cell(row['latitude'], row['longitude'], 5), axis=1)
        
        h3_dict = {}
        for h3_index, group in filtered_data.groupby('h3_index'):
            coords = h3.cell_to_boundary(h3_index)
            total_count = len(group)
            normal_count = sum(group['연결상태'] == 'normal')
            disc_count = sum(group['연결상태'] == 'disc.')

            h3_dict[h3_index] = {
                'coords': coords,
                'total_count': total_count,
                'normal_count': normal_count,
                'disc_count': disc_count,
            }

        # H3 경계 및 센서 개수 지도에 표시
        for h3_index, cell_data in h3_dict.items():
            boundary_coords = [(lat, lon) for lat, lon in cell_data['coords']]
            folium.PolyLine(boundary_coords, color="darkred", weight=4, opacity=0.8).add_to(m)

            # 셀 중심에 센서 수 표시
            lat_center = sum([lat for lat, lon in boundary_coords]) / len(boundary_coords)
            lon_center = sum([lon for lat, lon in boundary_coords]) / len(boundary_coords)
            label = f"{cell_data['total_count']} ({cell_data['normal_count']}/{cell_data['disc_count']})"
            folium.Marker(
                location=[lat_center, lon_center],
                popup=f"센서 수: {label}",
                icon=folium.DivIcon(html=f"""
                    <div style="text-align: center;">
                        <span style="font-size: 14pt; font-weight: bold; color: darkblue">{label}</span>
                    </div>
                """)
            ).add_to(m)

        # 개별 센서 위치 및 상태 표시
        marker_cluster = MarkerCluster().add_to(m)
        for idx, row in filtered_data.iterrows():
            color = "green" if row['연결상태'] == "normal" else "red"
            folium.Marker(
                location=[row['latitude'], row['longitude']],
                popup=f"단말번호: {row['단말번호']}\n연결상태: {row['연결상태']}",
                icon=folium.Icon(color=color)
            ).add_to(marker_cluster)

        # Streamlit에 Folium 지도 표시
        st_folium(m, width=700, height=500)
    else:
        st.write("위치 정보가 포함된 필터링된 데이터가 없습니다.")

def download_image_from_dropbox(url):
    # Dropbox 미리보기 링크를 다운로드 링크로 변경
    # Replace dl=0 with raw=1 to get image only
    download_url = url.replace("dl=0", "raw=1")
    response = requests.get(download_url)
    if response.status_code == 200:
        return BytesIO(response.content)  # 이미지 데이터를 BytesIO로 반환
    else:
        return None  # 이미지 다운로드 실패 시 None 반환


# 필터링된 데이터를 Word로 저장하는 버튼 추가
if st.button("필터링된 데이터를 MS Word로 저장"):
    # Word Document 객체 생성
    doc = Document()
    doc.add_heading("필터링된 MEMS 센서 데이터", 0)

    for idx, row in filtered_data.iterrows():
        # 6열 x 10행 테이블 생성
        table = doc.add_table(rows=10, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 테이블 가운데 정렬

        # 중앙 정렬 함수 정의
        def center_align(cell):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 1행: 단말번호
        table.cell(0, 0).text = "단말번호"
        cell = table.cell(0, 1).merge(table.cell(0, 5))
        cell.text = str(row.get("단말번호", ""))
        center_align(table.cell(0, 0))
        center_align(cell)

        # 2행: 관측소 코드, 시설구분, 시설구분세부
        table.cell(1, 0).text = "관측소 코드"
        table.cell(1, 1).text = str(row.get("관측소코드", ""))
        table.cell(1, 2).text = "시설구분"
        table.cell(1, 3).text = str(row.get("시설구분", ""))
        table.cell(1, 4).text = "시설구분세부"
        table.cell(1, 5).text = str(row.get("시설구분세부", ""))
        for col in range(6):
            center_align(table.cell(1, col))

        # 3행: 제조사, 설치시점, 연결상태
        table.cell(2, 0).text = "제조사"
        table.cell(2, 1).text = str(row.get("제조사", ""))
        table.cell(2, 2).text = "설치시점"
        table.cell(2, 3).text = str(row.get("설치시점", ""))
        table.cell(2, 4).text = "연결상태"
        table.cell(2, 5).text = str(row.get("연결상태", ""))
        for col in range(6):
            center_align(table.cell(2, col))

        # 4행: 주소
        table.cell(3, 0).text = "주소"
        cell = table.cell(3, 1).merge(table.cell(3, 5))
        cell.text = str(row.get("주소", ""))
        center_align(table.cell(3, 0))
        center_align(cell)

        # 5행: 위도, 경도, 고도
        table.cell(4, 0).text = "위도"
        table.cell(4, 1).text = str(row.get("latitude", ""))
        table.cell(4, 2).text = "경도"
        table.cell(4, 3).text = str(row.get("longitude", ""))
        table.cell(4, 4).text = "고도"
        table.cell(4, 5).text = str(row.get("고도", ""))
        for col in range(6):
            center_align(table.cell(4, col))

        # 6행: 설치층, 전체층, 축보정
        table.cell(5, 0).text = "설치층"
        table.cell(5, 1).text = str(row.get("설치층수", ""))
        table.cell(5, 2).text = "전체층"
        table.cell(5, 3).text = str(row.get("건물전체층수", ""))
        table.cell(5, 4).text = "축보정"
        table.cell(5, 5).text = str(row.get("축보정", ""))
        for col in range(6):
            center_align(table.cell(5, col))

        # 7행: H3 Cell, 센서 품질, 통신 품질
        table.cell(6, 0).text = "H3 Cell"
        table.cell(6, 1).text = str(row.get("H3 Cell", ""))
        table.cell(6, 2).text = "센서 품질"
        table.cell(6, 3).text = str(row.get("센서 품질", ""))
        table.cell(6, 4).text = "통신 품질"
        table.cell(6, 5).text = str(row.get("통신 품질", ""))
        for col in range(6):
            center_align(table.cell(6, col))

        # 8행: H3 혼잡여부, 센서 교체 필요 여부, 통신 품질 안정 여부
        table.cell(7, 0).text = "H3 혼잡여부"
        table.cell(7, 1).text = str(row.get("H3_Category", ""))
        table.cell(7, 2).text = "센서 교체 필요 여부"
        table.cell(7, 3).text = str(row.get("Sensor_Replacement_Status", ""))
        table.cell(7, 4).text = "통신 품질 안정 여부"
        table.cell(7, 5).text = str(row.get("Communication_Quality_Status", ""))
        for col in range(6):
            center_align(table.cell(7, col))

        # 9행: 현장 설치 사진 링크
        table.cell(8, 0).text = "현장 설치 사진"
        cell = table.cell(8, 1).merge(table.cell(8, 5))
        cell.text = "사진 링크: " + str(row.get("현장 설치 사진", ""))
        center_align(table.cell(8, 0))
        center_align(cell)

        # 10행: 추가 이미지 링크 1, 2
        cell = table.cell(9, 0).merge(table.cell(9, 2))
        cell.text = str(row.get("Image_Link_1", "#Image_link1"))
        center_align(cell)

        cell = table.cell(9, 3).merge(table.cell(9, 5))
        cell.text = str(row.get("Image_Link_2", "#Image_link2"))
        center_align(cell)

        # 이미지 링크가 있는 경우 이미지 삽입
        image_url = row.get("Image_Link_1", None)
        # replace nan with None
        if image_url is not None and not pd.isna(image_url):
            image_data = download_image_from_dropbox(image_url)
            if image_data:
                doc.add_paragraph("현장 설치 사진:")
                doc.add_picture(image_data, width=Inches(4.5))  # 이미지 폭 조절
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph("이미지를 로드할 수 없습니다.")

        # 페이지 나누기
        doc.add_page_break()  # 새 단말번호 작성 시 새 페이지로 이동

    # Word 파일을 BytesIO 버퍼에 저장
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.success("데이터가 MS Word 파일로 저장되었습니다.")
    st.download_button("다운로드: 필터링된 데이터 Word 파일", data=buffer, file_name="filtered_sensor_data.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

